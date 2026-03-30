from docx import Document
from docx.shared import Pt, Inches, MM
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from config import Config

def export_to_word(resume_data, experience_level, filename):
    """Export resume to Word format"""
    
    doc = Document()
    
    # Set margins for A4
    sections = doc.sections
    for section in sections:
        section.top_margin = MM(10)
        section.bottom_margin = MM(10)
        section.left_margin = MM(10)
        section.right_margin = MM(10)
    
    # Title (Name)
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(resume_data['full_name'].upper())
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact_para = doc.add_paragraph()
    contact_text = f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}"
    if resume_data.get('linkedin'):
        contact_text += f" | {resume_data['linkedin']}"
    if resume_data.get('portfolio'):
        contact_text += f" | {resume_data['portfolio']}"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=2)
    doc.add_paragraph(resume_data['professional_summary'])
    
    # Core Competencies
    doc.add_heading('CORE COMPETENCIES', level=2)
    competencies = ', '.join(resume_data.get('core_competencies', []))
    doc.add_paragraph(competencies)
    
    # Technical Expertise
    if resume_data.get('technical_expertise'):
        doc.add_heading('TECHNICAL EXPERTISE', level=2)
        tech = ', '.join(resume_data['technical_expertise'])
        doc.add_paragraph(tech)
    
    # Functional Expertise
    if resume_data.get('functional_expertise'):
        doc.add_heading('FUNCTIONAL EXPERTISE', level=2)
        func = ', '.join(resume_data['functional_expertise'])
        doc.add_paragraph(func)
    
    # Professional Experience
    if resume_data.get('experience'):
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
        for exp in resume_data['experience']:
            doc.add_heading(f"{exp['role']} | {exp['organization_name']}", level=3)
            doc.add_paragraph(f"{exp['job_start_year']} - {exp['job_end_year']}")
            
            for bullet in exp['project_detail']:
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style='List Bullet')
    
    # Education
    doc.add_heading('EDUCATION', level=2)
    if resume_data.get('graduate_degree'):
        grad = resume_data['graduate_degree']
        edu_text = f"{grad['degree_name']} | {grad['institution_name']} | {grad['graduation_year']}"
        doc.add_paragraph(edu_text)
    
    if resume_data.get('post_graduate_degree') and resume_data['post_graduate_degree'].get('degree_name'):
        post_grad = resume_data['post_graduate_degree']
        edu_text = f"{post_grad['degree_name']} | {post_grad['institution_name']} | {post_grad['graduation_year']}"
        doc.add_paragraph(edu_text)
    
    # Certifications
    if resume_data.get('certifications'):
        doc.add_heading('CERTIFICATIONS', level=2)
        for cert in resume_data['certifications']:
            cert_text = f"{cert['certification_name']} | {cert['institution_name']} | {cert['certification_year']}"
            doc.add_paragraph(cert_text)
    
    # Save Document
    os.makedirs(Config.EXPORT_DIR, exist_ok=True)
    word_path = os.path.join(Config.EXPORT_DIR, f"{filename}.docx")
    doc.save(word_path)
    
    return word_path
